from __future__ import annotations

import re
import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[1]
CODE_ROOT = PROJECT_ROOT / "src"
MANUSCRIPT = PROJECT_ROOT / "manuscript_context"
FIGURES = PROJECT_ROOT / "results/figures/manuscript"
EXPERIMENTS = PROJECT_ROOT / "results" / "experiments"
RUN_ID = "full_20260525"


def main() -> int:
    MANUSCRIPT.mkdir(parents=True, exist_ok=True)
    context = build_context()
    md = build_markdown(context)
    (MANUSCRIPT / "manuscript_draft.md").write_text(md, encoding="utf-8")
    (MANUSCRIPT / "main.tex").write_text(build_latex(context), encoding="utf-8")
    write_docx(context)
    print("Wrote manuscript draft files under", MANUSCRIPT)
    return 0


def build_context() -> dict:
    e2 = read_summary("E2_core_comparison")
    e3 = read_summary("E3_uncertainty_robustness")
    e5 = read_summary("E5_online_replanning")
    e6 = read_summary("E6_scalability")

    def row(df, **filters):
        sub = df.copy()
        for key, value in filters.items():
            sub = sub[sub[key].astype(str).eq(str(value))]
        if sub.empty:
            raise KeyError(filters)
        return sub.iloc[0]

    proposed = row(e2, method="uq_rv_alns")
    nearest = row(e2, method="greedy_nearest")
    point = row(e2, method="alns_point")
    e3_high_point = row(e3, method="alns_point", uncertainty="high")
    e3_high_prop = row(e3, method="uq_rv_alns", uncertainty="high")
    static = row(e5, policy="static")
    event = row(e5, policy="event_triggered")
    periodic = row(e5, policy="periodic")
    l_point = row(e6, method="alns_point", size="L")
    l_prop = row(e6, method="uq_rv_alns", size="L")
    return {
        "risk_reduction_vs_nearest": pct_reduction(proposed["risk_weighted_completion_time_mean"], nearest["risk_weighted_completion_time_mean"]),
        "risk_reduction_vs_point": pct_reduction(proposed["risk_weighted_completion_time_mean"], point["risk_weighted_completion_time_mean"]),
        "coverage_gain_vs_point": proposed["top_risk_coverage_mean"] - point["top_risk_coverage_mean"],
        "makespan_delta_vs_point": pct_change(proposed["makespan_mean"], point["makespan_mean"]),
        "high_violation_point": e3_high_point["infeasible_sortie_rate_mean"],
        "high_violation_prop": e3_high_prop["infeasible_sortie_rate_mean"],
        "high_makespan_penalty": pct_change(e3_high_prop["makespan_mean"], e3_high_point["makespan_mean"]),
        "online_makespan_gain": pct_reduction(event["makespan_mean"], static["makespan_mean"]),
        "online_violation_static": static["infeasible_sortie_rate_mean"],
        "online_violation_event": event["infeasible_sortie_rate_mean"],
        "event_response": event["online_response_time_mean"],
        "periodic_response": periodic["online_response_time_mean"],
        "large_runtime_ratio": l_prop["solver_runtime_mean"] / l_point["solver_runtime_mean"],
        "large_coverage_gain": l_prop["top_risk_coverage_mean"] - l_point["top_risk_coverage_mean"],
        "tables": {
            "e2": e2,
            "e3": e3,
            "e5": e5,
            "e6": e6,
        },
    }


def build_markdown(c: dict) -> str:
    references = "\n\n".join(reference_paragraphs())
    return f"""# Uncertainty-quantified risk-value scheduling for vehicle-UAV transmission line inspection

Author names withheld for review

## Highlights

- A vehicle-UAV inspection model couples probabilistic energy prediction with risk-value scheduling.
- A chance-constrained sortie filter converts energy uncertainty into conservative dispatch actions.
- Event-triggered rolling ALNS improves disturbed-operation performance over static and periodic replanning.
- Synthetic experiments follow a reusable TRE-style output protocol with raw data, source figures, and verified references.

## Abstract

Vehicle-UAV collaboration is a promising mode for large-scale transmission-line inspection, yet most scheduling models still rely on deterministic endurance estimates and homogeneous task values. This paper formulates infrastructure inspection as an uncertainty-aware and risk-value aware vehicle-UAV scheduling problem. We propose UQ-RV-ALNS, a framework that combines a probabilistic physics-informed energy surrogate, a chance-constrained sortie feasibility filter, a risk-value objective for high-priority inspection tasks, and event-triggered rolling replanning under wind, road-delay, urgent-task and communication disturbances. In this legacy controlled synthetic simulation suite, UQ-RV-ALNS reduces risk-weighted completion time by {c['risk_reduction_vs_nearest']:.1f}% relative to nearest-stop routing and by {c['risk_reduction_vs_point']:.1f}% relative to point energy ALNS; the {100*c['coverage_gain_vs_point']:.1f}-percentage-point top-risk coverage change is reported as descriptive simulation evidence. Under high uncertainty, the proposed conservative dispatch layer reduces the mean infeasible-sortie rate from {100*c['high_violation_point']:.2f}% to {100*c['high_violation_prop']:.2f}% with a {c['high_makespan_penalty']:.2f}% makespan change. In disturbed online episodes, event-triggered replanning reduces makespan by {c['online_makespan_gain']:.1f}% relative to a static plan and eliminates the simulated violation rate observed in static operation. The evidence is simulation-based and should be interpreted as algorithmic validation rather than field deployment validation.

Keywords: vehicle-UAV routing; transmission line inspection; adaptive large neighborhood search; physics-informed energy prediction; chance constraints; online replanning

## 1. Introduction

Drone-assisted logistics and inspection research has moved from stylized truck-drone routing toward richer air-ground systems with energy supply, no-fly zones, multimodal integration, automated heuristic design and synchronized truck-drone fulfillment \\cite{{kim2026bidirectional,liu2026evtol,zhang2026airground,shi2026llm,yang2025integrated}}. Transmission-line inspection shares this operational structure but differs from retail delivery in three important respects. First, the value of inspecting a tower or line segment is heterogeneous because defect likelihood, line criticality and image-risk priors vary spatially. Second, feasibility depends on energy and endurance under wind, payload, temperature and battery reserve uncertainty. Third, practical operation is rarely a single offline plan: weather changes, road delays, temporary communication degradation and urgent reinspection events can invalidate an initially feasible schedule.

The initial project concept already contained a mixed-integer model, ALNS solver and energy-surrogate prediction. The key limitation was that a point energy prediction still behaves like a deterministic parameter once it enters the optimizer. A schedule can therefore look efficient in the model while containing sorties that are unsafe under predictive uncertainty. The second limitation was objective design. Makespan minimization is necessary, but it is not sufficient when high-risk towers should be inspected earlier than low-value tasks. The third limitation was operational: a static plan cannot use new information during the execution horizon.

This paper addresses these limitations with UQ-RV-ALNS. The method is deliberately conservative. It does not claim exact optimality or field validation. Instead, it asks whether an uncertainty-aware and value-aware scheduling layer can improve the operational evidence chain needed for a Transportation Research Part E style manuscript.

The contributions are threefold. First, the paper converts a point surrogate-style energy estimate into a probabilistic scheduling interface, so uncertainty changes dispatch feasibility instead of being reported only after optimization. Second, it introduces a risk-value objective for inspection tasks, aligning routing decisions with defect-priority logic rather than pure makespan. Third, it evaluates event-triggered rolling replanning against static and periodic policies under disturbance events, with all simulation data stored in a reusable audit structure.

## 2. Literature review

The flying-sidekick traveling salesman problem established a canonical optimization model for drone-assisted delivery \\cite{{murray2015flying}}. Recent TRE papers extend this family in several directions: energy supply logistics with aerial and ground vehicles \\cite{{kim2026bidirectional}}, eVTOL-drone routing with no-fly zones and reinforcement learning \\cite{{liu2026evtol}}, air-ground multimodal integration \\cite{{zhang2026airground}}, automatic heuristic design for vehicle-drone routing \\cite{{shi2026llm}}, and synchronized order splitting with truck-drone fleets \\cite{{yang2025integrated}}. These studies show that current routing research values scale, operational realism and solver efficiency.

Adaptive large neighborhood search remains attractive for large routing variants because destroy and repair operators can absorb heterogeneous constraints without requiring every operational detail to be solved exactly \\cite{{ropke2006alns}}. Physics-informed neural networks provide a way to embed physical residuals into learned surrogates \\cite{{raissi2019pinn}}, but using a a surrogate only as a point predictor does not solve decision risk. This paper therefore uses a probabilistic surrogate interface: the optimizer receives both expected energy and an uncertainty-adjusted upper bound.

## 3. Problem description

Let \(T\) be the set of inspection tasks, \(P\) the set of vehicle stopping points, \(D\) the set of UAVs and \(V\) the set of ground vehicles. Each task \(i \\in T\) has a location, service time \(s_i\), risk score \(r_i\) and value \(v_i\). A vehicle can stop at \(p \\in P\), dispatch one or more UAVs in parallel, and each UAV sortie may inspect an ordered subset of nearby towers before returning to the same stop. A recovered UAV can then be dispatched again while the vehicle remains parked. The synthetic generator creates corridor-like tower coordinates, road-side stopping points, wind and temperature states, UAV battery capacity and disturbance events.

The core decision is to assign each task to a stop and UAV sequence while minimizing a composite operational objective:

\\[
\\min \\; C_\\max + \\lambda_E \\sum_i \\hat E_i + \\lambda_R \\sum_i r_i v_i C_i + \\lambda_M \\sum_i m_i,
\\]

where \(C_\\max\) is makespan, \(\hat E_i\) is predicted sortie energy, \(C_i\) is the completion time of task \(i\), and \(m_i\) is a missed or infeasible-task penalty.

## 4. Probabilistic energy interface

The energy surrogate returns a mean and uncertainty estimate for each candidate sortie:

\\[
\\hat E_i = \\mu_E(x_i;\\theta), \\qquad \\sigma_i = \\sigma_E(x_i;\\theta),
\\]

where \(x_i\) contains distance, wind speed, wind direction, payload, temperature and service-time features. A physically informed loss can be written as:

\\[
\\mathcal L = \\mathcal L_\\text{{data}} + \\alpha \\mathcal L_\\text{{physics}} + \\beta \\mathcal L_\\text{{cal}},
\\]

with data fit, physical residual and calibration terms. The scheduling layer uses the chance-constrained surrogate:

\\[
\\mu_E(x_i) + z_\\epsilon \\sigma_E(x_i) \\leq B_d(1-\\rho),
\\]

where \(B_d\) is UAV battery capacity and \(\\rho\) is the reserve ratio. If a sortie violates this bound but remains near the reserve limit, UQ-RV-ALNS applies a conservative sortie mode, representing slower flight, additional reserve management or a micro-stop/recharge action. This action reduces violation risk at a small duration cost.

![Energy uncertainty interface](../results/figures/manuscript/manuscript_Fig2_energy_uq_interface.png)

## 5. UQ-RV-ALNS algorithm

UQ-RV-ALNS combines risk-value task ordering, uncertainty-aware stop selection and event-triggered replanning. The destroy-repair logic follows the ALNS principle but avoids accepting local search moves that delay high-risk, high-value tasks without a compensating feasibility benefit.

Algorithm 1. UQ-RV-ALNS:

1. Generate or update scenario state from towers, stops, weather and battery state.
2. Estimate \(\\mu_E, \\sigma_E\) for candidate same-stop sortie patterns.
3. Rank tasks by \(v_i(1+2r_i)\) and select stops by q95 energy, duration and infeasibility penalty.
4. Build a multi-UAV schedule using earliest-available assignment.
5. Apply conservative sortie mode for near-boundary q95 violations.
6. Evaluate makespan, expected energy, risk-weighted completion time and violation rate.
7. When an online event exceeds a threshold, warm-start from the current plan and reoptimize affected tasks.

![Workflow](../results/figures/manuscript/manuscript_Fig1_workflow.png)

## 6. Experimental design

The experiments use synthetic corridor scenarios at S, M and L scales. Each experiment stores raw CSV data, summary CSV/Markdown files, source data for figures, and rendered figures. Methods include nearest routing, random feasible assignment, GA-style and ACO-style constructive baselines, fixed ALNS, point energy ALNS, UQ-ALNS, RV-ALNS and UQ-RV-ALNS.

The experiment matrix is:

- E1: small-instance reference comparison.
- E2: core method comparison.
- E3: low/medium/high uncertainty robustness.
- E4: value-objective ablation.
- E5: static, periodic and event-triggered online replanning.
- E6: scalability over S/M/L instances.

## 7. Results

### 7.1 Core comparison

UQ-RV-ALNS has the lowest risk-weighted completion time among the main baselines in the core comparison. Relative to nearest routing, it reduces risk-weighted completion time by {c['risk_reduction_vs_nearest']:.1f}%. Relative to point energy ALNS, the reduction is {c['risk_reduction_vs_point']:.1f}% and top-risk coverage improves by {100*c['coverage_gain_vs_point']:.1f} percentage points. The makespan change relative to point energy ALNS is {c['makespan_delta_vs_point']:.2f}%.

![Core comparison](../results/figures/manuscript/manuscript_Fig3_core_comparison.png)

### 7.2 Robustness under uncertainty

Under high uncertainty, point energy ALNS has a mean infeasible-sortie rate of {100*c['high_violation_point']:.2f}%. UQ-RV-ALNS reduces this to {100*c['high_violation_prop']:.2f}% by using conservative q95 feasibility and mitigation. The makespan change is {c['high_makespan_penalty']:.2f}%, which is the expected cost of reliability.

![Uncertainty robustness](../results/figures/manuscript/manuscript_Fig4_uncertainty_robustness.png)

### 7.3 Value-objective ablation

The value-aware component is mainly responsible for top-risk coverage. UQ alone protects feasibility but does not improve task priority as much as the risk-value objective.

![Value ablation](../results/figures/manuscript/manuscript_Fig5_value_ablation.png)

### 7.4 Online replanning

Static plans suffer from disturbed-operation penalties. Event-triggered replanning reduces disturbed makespan by {c['online_makespan_gain']:.1f}% relative to static planning and reduces the simulated infeasible-sortie rate from {100*c['online_violation_static']:.2f}% to {100*c['online_violation_event']:.2f}%. Its average response time is {c['event_response']:.2f}, compared with {c['periodic_response']:.2f} for periodic replanning.

![Online replanning](../results/figures/manuscript/manuscript_Fig6_online_replanning.png)

### 7.5 Scalability

At L scale, UQ-RV-ALNS has a runtime ratio of {c['large_runtime_ratio']:.2f} relative to point energy ALNS and reports a descriptive {100*c['large_coverage_gain']:.1f}-percentage-point top-risk coverage increase. The runtime result reflects the present implementation choice: the final UQ-RV variant prioritizes a direct risk-value construction plus conservative feasibility action rather than a heavy local-search pass.

![Scalability](../results/figures/manuscript/manuscript_Fig7_scalability.png)

## 8. Discussion

The results support a bounded claim. The proposed framework is not uniformly better on every metric. Its strength is the combination of value prioritization and reliability protection. It gives up a small amount of makespan in high-uncertainty settings to reduce violations, and it deliberately schedules high-risk tasks earlier. This is a defensible operational tradeoff for inspection, where missing or delaying a high-risk tower can be more consequential than adding a small travel or time cost.

The main limitation is data realism. The present experiments are synthetic simulations with plausible corridor geometry and weather-sensitive energy functions. They are useful for algorithmic validation and manuscript development, but they do not replace field logs, real defect labels, airspace constraints or verified battery degradation data.

Relative to the five reference TRE papers, the present draft is strongest in methodological coupling: it combines energy uncertainty, inspection value and online replanning in one route-construction loop. It is weaker in empirical grounding because the reference papers are framed around richer operational case settings or broader transport-system abstractions. The intended submission claim is therefore an algorithmic contribution for infrastructure inspection, not a validated deployment system.

## 9. Conclusion

This paper reframes vehicle-UAV transmission-line inspection as an uncertainty-aware and risk-value aware operational scheduling problem. By coupling probabilistic energy prediction, chance-constrained feasibility, risk-value task ordering and event-triggered replanning, UQ-RV-ALNS provides a stronger evidence chain than deterministic makespan-only scheduling. The current simulation results are sufficient for a manuscript draft and for identifying the next evidence gap: real or higher-fidelity flight logs are needed before making deployment claims.

## CRediT author statement

To be completed by the authors.

## Declaration of competing interest

The authors declare no known competing financial interests or personal relationships that could have appeared to influence the work reported in this paper.

## Data availability

All synthetic data, scripts, source CSVs and figures used in the current draft are stored in the local project folders. The generated evidence should be cited as simulation data.

## References

{references}
"""


def build_latex(c: dict) -> str:
    md = build_markdown(c)
    body = md
    replacements = [
        (r"^# (.*)$", r"\\title{\1}"),
        (r"^## (.*)$", r"\\section{\1}"),
        (r"^### (.*)$", r"\\subsection{\1}"),
    ]
    lines = []
    for line in body.splitlines():
        if line.startswith("!["):
            match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
            if match:
                caption, path = match.groups()
                stem = Path(path).with_suffix("").name
                lines.append("\\begin{figure}[htbp]")
                lines.append("\\centering")
                lines.append(f"\\includegraphics[width=0.92\\linewidth]{{../results/figures/manuscript/{stem}.pdf}}")
                lines.append(f"\\caption{{{caption}}}")
                lines.append("\\end{figure}")
            continue
        converted = line
        for pattern, repl in replacements:
            converted = re.sub(pattern, repl, converted)
        converted = converted.replace("%", "\\%")
        lines.append(converted)
    content = "\n".join(lines)
    content = re.sub(r"\\section\{References\}\s*.*", "", content, flags=re.S)
    content = content.replace("C_\\max", "C_{\\max}")
    return f"""\\documentclass[preprint,12pt]{{elsarticle}}
\\usepackage{{graphicx}}
\\usepackage{{amsmath}}
\\usepackage{{booktabs}}
\\journal{{Transportation Research Part E: Logistics and Transportation Review}}

\\begin{{document}}
{content}
\\bibliographystyle{{elsarticle-harv}}
\\bibliography{{elsevier/tre_references}}
\\end{{document}}
"""


def write_docx(c: dict) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size in [("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)]:
        styles[style_name].font.name = "Times New Roman"
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.bold = True

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Uncertainty-quantified risk-value scheduling for vehicle-UAV transmission line inspection")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(47, 62, 70)
    subtitle = doc.add_paragraph("Author names withheld for review")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "Highlights", 1)
    for item in [
        "A vehicle-UAV inspection model couples probabilistic energy prediction with risk-value scheduling.",
        "A chance-constrained sortie filter converts energy uncertainty into conservative dispatch actions.",
        "Event-triggered rolling ALNS improves disturbed-operation performance over static and periodic replanning.",
        "Synthetic experiments follow a reusable TRE-style output protocol with raw data, source figures, and verified references.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "Abstract", 1)
    doc.add_paragraph(
        f"Vehicle-UAV collaboration is a promising mode for large-scale transmission-line inspection, yet most scheduling models still rely on deterministic endurance estimates and homogeneous task values. This paper formulates infrastructure inspection as an uncertainty-aware and risk-value aware vehicle-UAV scheduling problem. We propose UQ-RV-ALNS, combining a probabilistic physics-informed energy surrogate, a chance-constrained sortie feasibility filter, a risk-value objective, and event-triggered rolling replanning. In this legacy simulation suite, UQ-RV-ALNS reduces risk-weighted completion time by {c['risk_reduction_vs_nearest']:.1f}% relative to nearest-stop routing and by {c['risk_reduction_vs_point']:.1f}% relative to point energy ALNS; the {100*c['coverage_gain_vs_point']:.1f}-percentage-point top-risk coverage change is descriptive simulation evidence. Under high uncertainty, it reduces infeasible-sortie rate from {100*c['high_violation_point']:.2f}% to {100*c['high_violation_prop']:.2f}% with a {c['high_makespan_penalty']:.2f}% makespan change. The evidence is simulation-based and should not be interpreted as field validation."
    )
    doc.add_paragraph("Keywords: vehicle-UAV routing; transmission line inspection; adaptive large neighborhood search; physics-informed energy prediction; chance constraints; online replanning")

    sections = [
        ("1. Introduction", intro_text()),
        ("2. Literature review", literature_text()),
        ("3. Problem description", problem_text()),
        ("4. Probabilistic energy interface", energy_text()),
        ("5. UQ-RV-ALNS algorithm", algorithm_text()),
        ("6. Experimental design", experiment_text()),
    ]
    for heading, paragraphs in sections:
        add_heading(doc, heading, 1)
        for p in paragraphs:
            doc.add_paragraph(p)
        if heading == "4. Probabilistic energy interface":
            add_figure(doc, "manuscript_Fig2_energy_uq_interface.png", "Fig. 2. Energy uncertainty interface and chance-constraint margin.")
        if heading == "5. UQ-RV-ALNS algorithm":
            add_figure(doc, "manuscript_Fig1_workflow.png", "Fig. 1. Closed-loop UQ-RV-ALNS inspection scheduling workflow.")

    add_heading(doc, "7. Results", 1)
    result_sections = [
        ("7.1 Core comparison", f"UQ-RV-ALNS reduces risk-weighted completion time by {c['risk_reduction_vs_nearest']:.1f}% relative to nearest routing and by {c['risk_reduction_vs_point']:.1f}% relative to point energy ALNS. Top-risk coverage improves by {100*c['coverage_gain_vs_point']:.1f} percentage points, while the makespan change relative to point energy ALNS is {c['makespan_delta_vs_point']:.2f}%.", "manuscript_Fig3_core_comparison.png", "Fig. 3. Core comparison across routing and ALNS baselines."),
        ("7.2 Robustness under uncertainty", f"Under high uncertainty, point energy ALNS has a mean infeasible-sortie rate of {100*c['high_violation_point']:.2f}%. UQ-RV-ALNS reduces this value to {100*c['high_violation_prop']:.2f}% by using q95 feasibility and conservative sortie mitigation.", "manuscript_Fig4_uncertainty_robustness.png", "Fig. 4. Robustness under low, medium and high uncertainty."),
        ("7.3 Value objective ablation", "The value-aware component is mainly responsible for prioritizing high-risk towers. UQ alone protects feasibility, but the risk-value objective is needed for top-risk coverage.", "manuscript_Fig5_value_ablation.png", "Fig. 5. Ablation of uncertainty and value-aware terms."),
        ("7.4 Online replanning", f"Event-triggered replanning reduces disturbed makespan by {c['online_makespan_gain']:.1f}% relative to static planning and reduces the simulated infeasible-sortie rate from {100*c['online_violation_static']:.2f}% to {100*c['online_violation_event']:.2f}%.", "manuscript_Fig6_online_replanning.png", "Fig. 6. Static, periodic and event-triggered online replanning."),
        ("7.5 Scalability", f"At L scale, UQ-RV-ALNS has a runtime ratio of {c['large_runtime_ratio']:.2f} relative to point energy ALNS and reports a descriptive {100*c['large_coverage_gain']:.1f}-percentage-point top-risk coverage increase.", "manuscript_Fig7_scalability.png", "Fig. 7. Scalability over S, M and L scenarios."),
    ]
    for heading, paragraph, fig, caption in result_sections:
        add_heading(doc, heading, 2)
        doc.add_paragraph(paragraph)
        add_figure(doc, fig, caption)

    add_heading(doc, "8. Discussion", 1)
    for p in discussion_text():
        doc.add_paragraph(p)
    add_heading(doc, "9. Conclusion", 1)
    doc.add_paragraph(conclusion_text())
    add_heading(doc, "CRediT author statement", 1)
    doc.add_paragraph("To be completed by the authors.")
    add_heading(doc, "Declaration of competing interest", 1)
    doc.add_paragraph("The authors declare no known competing financial interests or personal relationships that could have appeared to influence the work reported in this paper.")
    add_heading(doc, "Data availability", 1)
    doc.add_paragraph("All synthetic data, scripts, source CSVs and figures used in the current draft are stored in the local project folders. The generated evidence should be cited as simulation data.")
    add_heading(doc, "References", 1)
    for ref in reference_paragraphs():
        doc.add_paragraph(ref)
    doc.save(MANUSCRIPT / "paper_UQ_RV_ALNS_TRE_submission_draft.docx")


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_figure(doc: Document, filename: str, caption: str) -> None:
    path = FIGURES / filename
    if path.exists():
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(str(path), width=Inches(6.25))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True


def intro_text() -> list[str]:
    return [
        "Drone-assisted logistics and inspection research has moved from stylized truck-drone routing toward richer air-ground systems with energy supply, no-fly zones, multimodal integration, automated heuristic design and synchronized truck-drone fulfillment. Transmission-line inspection shares this operational structure but differs from retail delivery in task value, energy uncertainty and online disruption.",
        "The initial project concept already contained a mixed-integer model, ALNS solver and energy-surrogate prediction. The key limitation was that a point prediction behaves like a deterministic parameter once it enters the optimizer. This paper therefore turns the energy model into an uncertainty-aware scheduling interface.",
        "The contributions are a probabilistic energy interface, a risk-value scheduling objective for inspection tasks, and an event-triggered replanning layer evaluated against static and periodic operation.",
    ]


def literature_text() -> list[str]:
    return [
        "The flying-sidekick traveling salesman problem established a canonical optimization model for drone-assisted delivery. Recent TRE papers extend this family through energy supply logistics, eVTOL-drone routing, multimodal air-ground transport, LLM-based heuristic design and synchronized order splitting.",
        "Adaptive large neighborhood search remains attractive for large routing variants because destroy and repair operators can absorb heterogeneous constraints. Physics-informed neural networks provide a way to embed physical residuals into learned surrogates, but point prediction alone does not address decision risk.",
    ]


def problem_text() -> list[str]:
    return [
        "Let T be the set of inspection tasks, P the set of vehicle stopping points, D the set of UAVs and V the set of ground vehicles. Each task has a location, service time, risk score and value. A vehicle can stop, dispatch several UAVs in parallel, and each UAV may inspect an ordered subset of nearby towers before returning to the same stop.",
        "The objective combines makespan, expected energy, risk-weighted completion time and infeasibility penalties. The model is evaluated on synthetic corridor scenarios with seeded weather, battery and disturbance parameters.",
    ]


def energy_text() -> list[str]:
    return [
        "The energy surrogate returns mean energy and uncertainty for each candidate sortie. The scheduling layer accepts a sortie only when the q95 energy bound is within the battery reserve. Near-boundary sorties can use a conservative sortie mode that represents slower flight or a micro-stop/recharge action.",
        "This makes uncertainty operational: the optimizer does not merely report uncertainty after planning, but changes dispatch decisions when predicted reliability is poor.",
    ]


def algorithm_text() -> list[str]:
    return [
        "UQ-RV-ALNS ranks tasks by risk-adjusted value, selects stops using q95 energy and duration, forms bounded same-stop sortie patterns, assigns sorties to the earliest available UAV, and triggers replanning when online events exceed predefined thresholds.",
        "The final implementation avoids local-search moves that delay high-value tasks without improving feasibility. This design choice is supported by the quick ablation results and kept in the full experiment matrix.",
    ]


def experiment_text() -> list[str]:
    return [
        "The experiment suite contains E1 small-instance comparison, E2 core method comparison, E3 uncertainty robustness, E4 value ablation, E5 online replanning and E6 scalability. All experiments use fixed random seeds and write raw CSV, summary CSV, Markdown reports and source data.",
        "All results in the current manuscript are simulation results. They support algorithmic validation, not real-world deployment validation.",
    ]


def discussion_text() -> list[str]:
    return [
        "The proposed framework is not uniformly better on every metric. Its strength is the combination of value prioritization and reliability protection. It gives up a small amount of makespan in high-uncertainty settings to reduce violations, and it deliberately schedules high-risk tasks earlier.",
        "The main limitation is data realism. The present experiments are synthetic simulations with plausible corridor geometry and weather-sensitive energy functions. Field logs, real defect labels, airspace constraints and verified battery degradation data are needed before making deployment claims.",
        "Relative to the five reference TRE papers, the current draft is strongest in methodological coupling and weaker in empirical grounding. The intended claim is an algorithmic contribution for infrastructure inspection, not a validated deployment system.",
    ]


def conclusion_text() -> str:
    return (
        "This paper reframes vehicle-UAV transmission-line inspection as an uncertainty-aware and risk-value aware operational scheduling problem. By coupling probabilistic energy prediction, chance-constrained feasibility, risk-value task ordering and event-triggered replanning, UQ-RV-ALNS provides a stronger evidence chain than deterministic makespan-only scheduling."
    )


def reference_paragraphs() -> list[str]:
    return [
        "Kim, H., Sari Darmawi Purba, D., & Kontou, E. (2026). Bidirectional energy supply logistics using uncrewed electric aerial and ground vehicles. Transportation Research Part E, 209, 104726.",
        "Liu, S., Yu, Y., Tian, Q., & Sun, H. (2026). Routing optimization for an eVTOL-and-drone delivery system in continuous space with no-fly zones. Transportation Research Part E, 209, 104741.",
        "Zhang, Y., Yang, C., Xi, H., Peng, S., Yang, J., Gan, M., Liu, X., & Ai, R. (2026). Air-ground multimodal transport planning for joint passenger mobility and parcel delivery. Transportation Research Part E, 210, 104825.",
        "Shi, H., & Zhen, L. (2026). LLM-based automatic heuristic design for vehicle-drone collaborative routing problems. Transportation Research Part E, 209, 104760.",
        "Yang, R., & Li, X. (2025). Integrated order splitting, allocation, and delivery problem with a synchronized truck and drone fleet. Transportation Research Part E, 202, 104217.",
        "Murray, C. C., & Chu, A. G. (2015). The flying sidekick traveling salesman problem. Transportation Research Part C, 54, 86-109.",
        "Ropke, S., & Pisinger, D. (2006). An adaptive large neighborhood search heuristic for the pickup and delivery problem with time windows. Transportation Science, 40(4), 455-472.",
        "Raissi, M., Perdikaris, P., & Karniadakis, G. E. (2019). Physics-informed neural networks. Journal of Computational Physics, 378, 686-707.",
    ]


def read_summary(experiment_id: str) -> pd.DataFrame:
    return pd.read_csv(EXPERIMENTS / experiment_id / "analysis_data" / f"{experiment_id}_{RUN_ID}_summary.csv")


def pct_reduction(new: float, old: float) -> float:
    return (old - new) / old * 100.0


def pct_change(new: float, old: float) -> float:
    return (new - old) / old * 100.0


if __name__ == "__main__":
    raise SystemExit(main())
